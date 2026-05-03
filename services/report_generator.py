from __future__ import annotations

import base64
import io
import math
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from reportlab.lib.colors import HexColor, black, grey, white
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm, mm
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    Image as RLImage,
    NextPageTemplate,
    PageBreak,
    PageTemplate,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

TECH_COLORS: Dict[str, str] = {
    "GSM": "#00AA00",
    "WCDMA": "#0055FF",
    "LTE": "#FF8800",
    "NR": "#8800FF",
}

CORP_BLUE = "#003366"
CORP_GREY = "#666666"


def _safe_str(val: Any, fmt: Optional[str] = None) -> str:
    if val is None or (isinstance(val, float) and math.isnan(val)):
        return "N/D"
    try:
        if fmt:
            return fmt.format(float(val) if isinstance(val, (int, float, str)) else val)
        if isinstance(val, float):
            return f"{val:.2f}"
        return str(val)
    except (ValueError, TypeError):
        return str(val)


def _build_ascii_azimuth_map(sectors: List[Dict[str, Any]]) -> str:
    lines: List[str] = []
    for sec in sectors:
        az = float(sec.get("azimuth", sec.get("Azimute", 0))) % 360
        tech = str(sec.get("technology", sec.get("Tecnologia", "?"))).upper()
        radius = sec.get("radius_km", sec.get("Raio", "?"))
        radius_str = f"{float(radius):.1f}km" if not isinstance(radius, str) else radius

        if 0 <= az < 45:
            d = "N  ↑"
        elif 45 <= az < 135:
            d = "E  →"
        elif 135 <= az < 225:
            d = "S  ↓"
        elif 225 <= az < 315:
            d = "W  ←"
        else:
            d = "N  ↑"
        lines.append(f"  {tech:6s} Az={az:3.0f}° {d}  R={radius_str}")
    return "\n".join(lines)


# ======================================================================
# CHART FUNCTIONS
# ======================================================================

def create_technology_pie_chart(df_summary: Dict[str, Any]) -> bytes:
    """
    Gráfico de pizza mostrando distribuição percentual de setores por tecnologia.

    Args:
        df_summary: Dicionário com contagem por tecnologia, ex:
                    {"GSM": 5, "WCDMA": 3, "LTE": 12, "NR": 2}

    Returns:
        PNG como bytes (io.BytesIO).
    """
    if not df_summary:
        df_summary = {"GSM": 1, "LTE": 1}

    labels = list(df_summary.keys())
    values = list(df_summary.values())
    colors = [TECH_COLORS.get(lbl, "#CCCCCC") for lbl in labels]

    fig, ax = plt.subplots(figsize=(5, 4))
    wedges, texts, autotexts = ax.pie(
        values,
        labels=None,
        autopct="%1.1f%%",
        colors=colors,
        startangle=90,
        pctdistance=0.7,
        wedgeprops={"edgecolor": "white", "linewidth": 1.5},
    )
    for a in autotexts:
        a.set_fontsize(10)
        a.set_color("white")
        a.set_fontweight("bold")
    ax.legend(wedges, labels, title="Tecnologia", loc="center left",
              bbox_to_anchor=(1, 0.5), fontsize=9)
    ax.set_title("Distribuição de Setores por Tecnologia", fontsize=12, fontweight="bold",
                 color=CORP_BLUE)
    plt.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return buf.getvalue()


def create_coverage_chart(stations_data: List[Dict[str, Any]]) -> bytes:
    """
    Gráfico combinado: scatter das torres no mapa + histograma de raios de cobertura.

    Args:
        stations_data: Lista de estações com lat/lon e setores com radius_km.

    Returns:
        PNG como bytes (io.BytesIO).
    """
    if not stations_data:
        fig, ax = plt.subplots(figsize=(6, 3))
        ax.text(0.5, 0.5, "Sem dados disponíveis", ha="center", va="center")
        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
        plt.close(fig)
        buf.seek(0)
        return buf.getvalue()

    lats: List[float] = []
    lons: List[float] = []
    tech_labels: List[str] = []
    radii: List[float] = []

    for st in stations_data:
        lat = float(st.get("lat", st.get("latitude", 0)))
        lon = float(st.get("lon", st.get("longitude", 0)))
        for sec in st.get("sectors", []):
            lats.append(lat)
            lons.append(lon)
            tech_labels.append(str(sec.get("technology", sec.get("Tecnologia", "?"))).upper()[:6])
            r = sec.get("radius_km", 0)
            try:
                radii.append(float(r))
            except (ValueError, TypeError):
                radii.append(0)

    fig, (ax_map, ax_hist) = plt.subplots(1, 2, figsize=(10, 4),
                                            gridspec_kw={"width_ratios": [1.5, 1]})

    unique_techs = sorted(set(tech_labels))
    for tech in unique_techs:
        mask = [t == tech for t in tech_labels]
        c = TECH_COLORS.get(tech, "#CCCCCC")
        ax_map.scatter(
            [lons[i] for i, m in enumerate(mask) if m],
            [lats[i] for i, m in enumerate(mask) if m],
            c=c, label=tech, s=40, alpha=0.8, edgecolors="black", linewidth=0.3,
        )
    ax_map.set_xlabel("Longitude", fontsize=9)
    ax_map.set_ylabel("Latitude", fontsize=9)
    ax_map.set_title("Distribuição Geográfica das Torres", fontsize=10, fontweight="bold",
                     color=CORP_BLUE)
    ax_map.legend(fontsize=7, loc="best")
    ax_map.tick_params(labelsize=7)
    ax_map.grid(True, alpha=0.2)

    filtered = [r for r in radii if r > 0]
    if filtered:
        ax_hist.hist(filtered, bins=min(15, len(filtered)), color=CORP_BLUE, alpha=0.7,
                      edgecolor="white")
        ax_hist.axvline(x=np.mean(filtered), color="red", linestyle="--", linewidth=1.5,
                         label=f"Média={np.mean(filtered):.2f} km")
        ax_hist.legend(fontsize=7)
    ax_hist.set_xlabel("Raio de Cobertura (km)", fontsize=9)
    ax_hist.set_ylabel("Frequência", fontsize=9)
    ax_hist.set_title("Histograma de Raios de Cobertura", fontsize=10, fontweight="bold",
                      color=CORP_BLUE)
    ax_hist.tick_params(labelsize=7)
    ax_hist.grid(True, alpha=0.2)

    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return buf.getvalue()


# ======================================================================
# PDF REPORT
# ======================================================================

def _pdf_header_footer(canvas, doc):
    canvas.saveState()
    canvas.setFont("Helvetica", 7)
    canvas.setFillColor(HexColor(CORP_GREY))
    canvas.drawString(2 * cm, 1.2 * cm,
                      f"RF Tower System | Gerado em {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    canvas.drawRightString(A4[0] - 2 * cm, 1.2 * cm, f"Página {doc.page}")
    canvas.setStrokeColor(HexColor(CORP_GREY))
    canvas.setLineWidth(0.3)
    canvas.line(2 * cm, 1.5 * cm, A4[0] - 2 * cm, 1.5 * cm)
    canvas.restoreState()


def generate_pdf_report(
    stations_data: List[Dict[str, Any]],
    output_path: str,
    title: str = "Relatório de Cobertura RF",
) -> str:
    """
    Gera relatório PDF profissional com capa, sumário executivo e detalhamento
    por estação.

    Estrutura:
    - Capa (pág 1): título, data, estatísticas, logo RF Tower System
    - Sumário Executivo (pág 2): tabela resumo, gráfico pizza, gráfico barras
    - Detalhamento por Estação (páginas seguintes): tabelas de setores + mapa ASCII
    - Rodapé em todas as páginas

    Args:
        stations_data: Lista de dicts de estações com setores simulados.
        output_path: Caminho do arquivo PDF de saída.
        title: Título do relatório.

    Returns:
        Caminho do arquivo gerado.
    """
    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        rightMargin=2 * cm,
        leftMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2.2 * cm,
        title=title,
        author="RF Tower System",
    )

    elements: List[Any] = []
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle("CapaTitulo", parent=styles["Title"], fontSize=22,
                                  spaceAfter=8, alignment=TA_CENTER,
                                  textColor=HexColor(CORP_BLUE))
    subtitle_style = ParagraphStyle("CapaSubtitulo", parent=styles["Normal"], fontSize=12,
                                     spaceAfter=4, alignment=TA_CENTER,
                                     textColor=HexColor(CORP_GREY))
    h1_style = ParagraphStyle("H1", parent=styles["Heading1"], fontSize=15,
                               spaceAfter=10, spaceBefore=20,
                               textColor=HexColor(CORP_BLUE))
    h2_style = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=13,
                               spaceAfter=8, spaceBefore=16,
                               textColor=HexColor(CORP_BLUE))
    normal_style = ParagraphStyle("NormalText", parent=styles["Normal"], fontSize=10,
                                   spaceAfter=6, leading=14)
    small_style = ParagraphStyle("Small", parent=normal_style, fontSize=8,
                                  textColor=HexColor(CORP_GREY))

    # -------- Cover --------
    elements.append(Spacer(1, 80))
    elements.append(Paragraph("RF Tower System", ParagraphStyle(
        "LogoCapa", parent=title_style, fontSize=28, textColor=HexColor(CORP_BLUE))))
    elements.append(Spacer(1, 30))
    elements.append(Paragraph(title, title_style))
    elements.append(Spacer(1, 20))
    elements.append(Paragraph(
        f"Gerado em: {datetime.now().strftime('%d/%m/%Y às %H:%M:%S')}", subtitle_style))
    elements.append(Spacer(1, 40))

    total_stations = len(stations_data)
    total_sectors = sum(len(st.get("sectors", [])) for st in stations_data)
    techs: Dict[str, int] = {}
    all_radii: List[float] = []
    for st in stations_data:
        for sec in st.get("sectors", []):
            t = str(sec.get("technology", sec.get("Tecnologia", "?"))).upper()
            techs[t] = techs.get(t, 0) + 1
            try:
                all_radii.append(float(sec.get("radius_km", 0)))
            except (ValueError, TypeError):
                pass
    avg_radius = np.mean(all_radii) if all_radii else 0.0

    summary_data = [
        ["Indicador", "Valor"],
        ["Estações analisadas", str(total_stations)],
        ["Total de setores", str(total_sectors)],
        ["Tecnologias encontradas", str(len(techs))],
        ["Raio médio de cobertura", f"{avg_radius:.2f} km"],
    ]
    for t_name, t_cnt in sorted(techs.items()):
        summary_data.append([f"  └ {t_name}", f"{t_cnt} setor(es)"])

    sum_table = Table(summary_data, colWidths=[8 * cm, 6 * cm])
    sum_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), HexColor(CORP_BLUE)),
        ("TEXTCOLOR", (0, 0), (-1, 0), white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("GRID", (0, 0), (-1, -1), 0.5, grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [white, HexColor("#f0f4f8")]),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    elements.append(Paragraph("Sumário Executivo", h1_style))
    elements.append(sum_table)
    elements.append(Spacer(1, 16))

    if techs:
        pie_png = create_technology_pie_chart(techs)
        pie_img = RLImage(io.BytesIO(pie_png), width=10 * cm, height=8 * cm)
        bar_png = create_coverage_chart(stations_data)
        bar_img = RLImage(io.BytesIO(bar_png), width=16 * cm, height=6.5 * cm)

        combined = [[pie_img, bar_img]]
        combined_table = Table(combined, colWidths=[8 * cm, 8 * cm])
        combined_table.setStyle(TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (-1, -1), 2),
            ("RIGHTPADDING", (0, 0), (-1, -1), 2),
        ]))
        elements.append(combined_table)

    elements.append(PageBreak())

    # -------- Per-station detail --------
    elements.append(Paragraph("Detalhamento por Estação", h1_style))
    elements.append(Paragraph(
        f"Esta seção apresenta os parâmetros de cada setor das {total_stations} "
        f"estações analisadas.", normal_style))
    elements.append(Spacer(1, 8))

    for i, st in enumerate(stations_data):
        st_id = str(st.get("station_id", st.get("numero_estacao", i + 1)))
        operadora = str(st.get("operadora", st.get("Torre Estação", "")))
        endereco = str(st.get("endereco", st.get("EnderecoEstacao", "")))
        lat = _safe_str(st.get("lat", st.get("latitude", "")))
        lon = _safe_str(st.get("lon", st.get("longitude", "")))

        elements.append(Paragraph(
            f"#{st_id} — {endereco}", h2_style))
        elements.append(Paragraph(
            f"Operadora: {operadora}  |  Coordenadas: ({lat}, {lon})", normal_style))
        elements.append(Spacer(1, 4))

        sector_rows = [["Tecnologia", "Azimute", "FreqTx (MHz)", "Ganho (dBi)",
                         "Altura (m)", "Potência (W)", "Raio (km)"]]
        sectors = st.get("sectors", [])
        for sec in sectors:
            sector_rows.append([
                str(sec.get("technology", sec.get("Tecnologia", "?"))).upper(),
                _safe_str(sec.get("azimuth", sec.get("Azimute", "")), "{:.0f}"),
                _safe_str(sec.get("freq_mhz", sec.get("FreqTxMHz", "")), "{:.1f}"),
                _safe_str(sec.get("antenna_gain_dbi", sec.get("GanhoAntena", "")),
                          "{:.1f}"),
                _safe_str(sec.get("tx_height_m", sec.get("AlturaAntena", "")),
                          "{:.1f}"),
                _safe_str(sec.get("tx_power_watts",
                                   sec.get("PotenciaTransmissorWatts", "")), "{:.1f}"),
                _safe_str(sec.get("radius_km", ""), "{:.2f}"),
            ])

        st_radii = []
        for sec in sectors:
            try:
                st_radii.append(float(sec.get("radius_km", 0)))
            except (ValueError, TypeError):
                pass
        st_avg = np.mean(st_radii) if st_radii else 0

        sec_table = Table(sector_rows, colWidths=[2.2 * cm] * 7)
        sec_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), HexColor(CORP_BLUE)),
            ("TEXTCOLOR", (0, 0), (-1, 0), white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("GRID", (0, 0), (-1, -1), 0.4, grey),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [white, HexColor("#f0f4f8")]),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        elements.append(sec_table)
        elements.append(Spacer(1, 2))
        elements.append(Paragraph(
            f"Raio médio de cobertura desta estação: <b>{st_avg:.2f} km</b>",
            small_style))

        ascii_map = _build_ascii_azimuth_map(sectors)
        elements.append(Paragraph(
            f"<font face='Courier' size='7'>{ascii_map}</font>",
            ParagraphStyle("MonoSmall", parent=normal_style, fontSize=7, leading=9)
        ))
        elements.append(Spacer(1, 10))

    elements.append(Paragraph("--- Fim do Relatório ---", ParagraphStyle(
        "FooterEnd", parent=small_style, fontSize=8, alignment=TA_CENTER)))

    doc.build(elements, onFirstPage=_pdf_header_footer, onLaterPages=_pdf_header_footer)
    return output_path


# ======================================================================
# DOCX REPORT
# ======================================================================


def _set_cell_shading(cell, color_hex: str):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _format_docx_table(table, header_color: str = "003366"):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

    for j, cell in enumerate(table.rows[0].cells):
        _set_cell_shading(cell, header_color)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(9)

    for i, row in enumerate(table.rows[1:], 1):
        bg = "F0F4F8" if i % 2 == 0 else "FFFFFF"
        for cell in row.cells:
            _set_cell_shading(cell, bg)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)


def generate_docx_report(
    stations_data: List[Dict[str, Any]],
    output_path: str,
    title: str = "Relatório de Cobertura RF",
) -> str:
    """
    Gera relatório DOCX profissional com sumário executivo e detalhamento
    por estação.

    Args:
        stations_data: Lista de dicts de estações com setores simulados.
        output_path: Caminho do arquivo DOCX de saída.
        title: Título do relatório.

    Returns:
        Caminho do arquivo gerado.
    """
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    header = doc.sections[0].header
    hp = header.paragraphs[0]
    hp.text = title
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in hp.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp_run = fp.add_run()
    fp_run.font.size = Pt(8)
    fp_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    from docx.oxml import OxmlElement
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    fp_run._r.append(fld_char1)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fp_run._r.append(instr)
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    fp_run._r.append(fld_char2)

    # Title
    h_title = doc.add_heading(title, level=0)
    h_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h_title.runs:
        run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = date_p.add_run(f"Gerado em: {datetime.now().strftime('%d/%m/%Y às %H:%M:%S')}")
    dr.font.size = Pt(9)
    dr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph("")

    # Executive summary
    doc.add_heading("Resumo Executivo", level=1)
    total_stations = len(stations_data)
    total_sectors = sum(len(st.get("sectors", [])) for st in stations_data)
    techs: Dict[str, int] = {}
    all_radii: List[float] = []
    for st in stations_data:
        for sec in st.get("sectors", []):
            t = str(sec.get("technology", sec.get("Tecnologia", "?"))).upper()
            techs[t] = techs.get(t, 0) + 1
            try:
                all_radii.append(float(sec.get("radius_km", 0)))
            except (ValueError, TypeError):
                pass
    avg_radius = np.mean(all_radii) if all_radii else 0.0

    p = doc.add_paragraph()
    p.add_run(
        f"Este relatório apresenta os resultados da simulação de cobertura "
        f"de radiofrequência para {total_stations} estação(ões) totalizando "
        f"{total_sectors} setor(es). "
    ).font.size = Pt(11)
    p.add_run(
        f"Foram identificadas {len(techs)} tecnologia(s): "
        f"{', '.join(sorted(techs.keys()))}. "
        f"O raio médio de cobertura calculado é de {avg_radius:.2f} km."
    ).font.size = Pt(11)

    sum_rows = len(techs) + 6
    sum_table = doc.add_table(rows=sum_rows, cols=2)
    headers = [
        ("Estações analisadas", str(total_stations)),
        ("Total de setores", str(total_sectors)),
        ("Tecnologias encontradas", str(len(techs))),
        ("Raio médio de cobertura", f"{avg_radius:.2f} km"),
    ]
    for j, (lab, val) in enumerate(headers):
        sum_table.rows[j].cells[0].text = lab
        sum_table.rows[j].cells[1].text = val
    offset = len(headers)
    sum_table.rows[offset].cells[0].text = "Setores por tecnologia"
    sum_table.rows[offset].cells[0].paragraphs[0].runs[0].font.bold = True
    sum_table.rows[offset].cells[1].text = ""
    for j, (t_name, t_cnt) in enumerate(sorted(techs.items()), offset + 1):
        sum_table.rows[j].cells[0].text = f"  {t_name}"
        sum_table.rows[j].cells[1].text = str(t_cnt)
    _format_docx_table(sum_table)

    doc.add_paragraph("")

    if techs:
        pie_png = create_technology_pie_chart(techs)
        pie_buf = io.BytesIO(pie_png)
        doc.add_picture(pie_buf, width=Inches(3.5))
        pie_buf.close()

    doc.add_page_break()

    # Per-station detail
    doc.add_heading("Detalhamento por Estação", level=1)

    for i, st in enumerate(stations_data):
        st_id = str(st.get("station_id", st.get("numero_estacao", i + 1)))
        operadora = str(st.get("operadora", st.get("Torre Estação", "")))
        endereco = str(st.get("endereco", st.get("EnderecoEstacao", "")))
        lat = _safe_str(st.get("lat", st.get("latitude", "")))
        lon = _safe_str(st.get("lon", st.get("longitude", "")))

        h2 = doc.add_heading(f"#{st_id} — {endereco}", level=2)
        for run in h2.runs:
            run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

        dp = doc.add_paragraph()
        dp.add_run(f"Operadora: {operadora}").font.size = Pt(10)
        dp.add_run(f"  |  Coordenadas: ({lat}, {lon})").font.size = Pt(10)

        sectors = st.get("sectors", [])
        cols = ["Tecnologia", "Azimute", "FreqTx (MHz)", "Ganho (dBi)",
                "Altura (m)", "Potência (W)", "Raio (km)"]
        tbl = doc.add_table(rows=len(sectors) + 1, cols=len(cols))
        for j, col_name in enumerate(cols):
            tbl.rows[0].cells[j].text = col_name
        for r, sec in enumerate(sectors):
            vals = [
                str(sec.get("technology", sec.get("Tecnologia", "?"))).upper(),
                _safe_str(sec.get("azimuth", sec.get("Azimute", "")), "{:.0f}"),
                _safe_str(sec.get("freq_mhz", sec.get("FreqTxMHz", "")), "{:.1f}"),
                _safe_str(sec.get("antenna_gain_dbi", sec.get("GanhoAntena", "")),
                          "{:.1f}"),
                _safe_str(sec.get("tx_height_m", sec.get("AlturaAntena", "")),
                          "{:.1f}"),
                _safe_str(sec.get("tx_power_watts",
                                   sec.get("PotenciaTransmissorWatts", "")),
                          "{:.1f}"),
                _safe_str(sec.get("radius_km", ""), "{:.2f}"),
            ]
            for j, v in enumerate(vals):
                tbl.rows[r + 1].cells[j].text = v
        _format_docx_table(tbl)

        st_radii = []
        for sec in sectors:
            try:
                st_radii.append(float(sec.get("radius_km", 0)))
            except (ValueError, TypeError):
                pass
        st_avg = np.mean(st_radii) if st_radii else 0

        cp = doc.add_paragraph()
        cr = cp.add_run(f"Raio médio de cobertura: {st_avg:.2f} km")
        cr.font.size = Pt(9)
        cr.font.italic = True
        cr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        ascii_map = _build_ascii_azimuth_map(sectors)
        mp = doc.add_paragraph()
        mr = mp.add_run(f"Orientação dos setores:\n{ascii_map}")
        mr.font.name = "Courier New"
        mr.font.size = Pt(7)

        doc.add_paragraph("")

    # Final
    doc.add_paragraph("")
    fp2 = doc.add_paragraph()
    fp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr2 = fp2.add_run("--- Fim do Relatório ---")
    fr2.font.size = Pt(8)
    fr2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.save(output_path)
    return output_path
